import os
from textblob import TextBlob
from docx import Document  # Ensure this import statement is included

# Function to get sentiment
def get_sentiment(text):
    return TextBlob(text).sentiment.polarity

# Function to extract text from a .docx file
def extract_text_from_docx(file_path):
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

# Function to read keywords from a .docx file
def read_keywords_from_docx(file_path):
    doc = Document(file_path)
    keywords = []
    for para in doc.paragraphs:
        if para.text.strip():  # Ensure the line is not empty
            keywords.extend(para.text.split(','))
    return [keyword.strip() for keyword in keywords]

# Function to analyze content related to specified keywords
def analyze_content(text, keywords):
    content_lower = text.lower()
    sentences = text.split('.')
    for sentence in sentences:
        if any(keyword.lower() in sentence.lower() for keyword in keywords):
            clean_sentence = sentence.strip()
            sentiment = get_sentiment(clean_sentence)
            sentiment_label = "Neutral"
            if sentiment > 0.05:
                sentiment_label = "Positive"
            elif sentiment < -0.05:
                sentiment_label = "Negative"
            print(f"Sentence: {clean_sentence}")
            print(f"Sentiment: {sentiment_label} ({sentiment:.2f})")
            print()

# Load keywords from file
keywords_file_path = 'keywordsInvestments.docx'
keywords = read_keywords_from_docx(keywords_file_path)

# Load and analyze document text
content_file_path = 'btcSediment.docx'
document_text = extract_text_from_docx(content_file_path)
analyze_content(document_text, keywords)

# Optional: Reporting functionality (kept from your initial code snippet)
def generate_report(sentences, sentiments, keywords):
    pass  # Placeholder until the function is implemented
